import pandas as pd
from docx.api import Document
import argparse


parser = argparse.ArgumentParser(description="Use to extract tables from word document.")
parser.add_argument('--file', '--f', type=str, default="std", help='Target file')
parser.add_argument('--keywords', '--k', type=str, nargs='+', default='std', help="Key words used to filter header")
parser.add_argument('--limit', '--l', type=str, default='std', help="Using first limit lines to filter header")
parser.add_argument('--output', '--o', type=str, default='std', help="Output file")
args = parser.parse_args()


def table_count(file: str):
    """
    显示word文件中有多少表格
    """
    document = Document(file)
    print(f"该word文档中共有{len(document.tables)}个表格")

    
def seperate_header_data(target, keyWords, limit):
    """
    用于在一个list中找到哪些行是header，返回两个list，header和data，分别对应表头和数据
    
    :param target: 目标list
    :param keyWords: 用于判断是否是表头的关键词，如果一行中包含了其中任意一个关键词，即被视作是表头
    :param limit: 最多检查的行数，因为表头一般来说最多到3级，不会太多
    """
    headerIdx = 0
    
    for row in target:
        if limit == 0:
            break
        for key in keyWords:
            if key in row:
                headerIdx += 1
                break
        limit -= 1
    
    header = target[:headerIdx]
    data = target[headerIdx:]
    return header, data, headerIdx        
        
    
def extract_one_table(document, idx, keyWords, limit) -> pd.DataFrame:
    """
    根据编号提取一个表格
    
    :param document: 目标文档
    :param idx: 表格编号，从0开始计
    :param keyWords: 用于判断是否是表头的关键词，如果一行中包含了其中任意一个关键词，即被视作是表头
    :param limit: 最多检查的行数，因为表头一般来说最多到3级，不会太多
    """
    rows = []
    table = document.tables[idx]
    
    # 将word中的目标表格按行扫描，并保存在一个list中
    for i in table.rows:
        row = []
        for j in i.cells:
            row.append(j.text)
        rows.append(row)
        
    # 分离header和data，并创建dataframe
    header, data, headerIdx = seperate_header_data(rows, keyWords, limit)
    if len(header) != 0:
        result = pd.DataFrame(data, columns=header)
    else:
        result = pd.DataFrame(data)
    return result, headerIdx


def extract_all_table(file, keyWords, limit, targetFile, targetIdx='all'):
    """
    提取word中的所有目标表格，并按顺序保存进excel文件中
    
    :param file: 目标文件路径
    :param keyWords: 用于判断是否是表头的关键词，如果一行中包含了其中任意一个关键词，即被视作是表头
    :param limit: 最多检查的行数，因为表头一般来说最多到3级，不会太多
    :param targetFile: 目标excel文件路径
    :param targetIdx: 目标表格的编号，从0开始计，默认为提取全部表格，可以用list指定编号
    """
    # 初始化文档对象和excel写入器
    document = Document(file)
    writer = pd.ExcelWriter(targetFile, engine='xlsxwriter')
    sheetIdx = 1 # 为了最终和word中的表标题对应，这里的sheet名字从1开始计
    print(f"目标文档中共发现{len(document.tables)}张表格")
    
    if targetIdx == 'all':
        targetIdx = list(range(len(document.tables)))
    elif not isinstance(targetIdx, list):
        targetIdx = [int(targetIdx)]
    elif isinstance(targetIdx, list):
        targetIdx = targetIdx
    
    # 按顺序提取文档中的表格，并保存df到目标excel中
    with pd.ExcelWriter(targetFile, engine='xlsxwriter') as writer:
        for idx in targetIdx:
            df, headerIdx = extract_one_table(document, idx, keyWords, limit)
            # df = pd.DataFrame(df[1:], columns=df[0])  # Use the first row as column names
            df.to_excel(writer, sheet_name=f"表{sheetIdx}")
            if headerIdx >= 1:
                writer.sheets[f"表{sheetIdx}"].set_row(headerIdx, None, None, {'hidden': True}) # 用于隐藏header下面多出来的一行空白行
            sheetIdx += 1

    
if __name__ == '__main__':
    # 目标文件
    if args.file == 'std':
        file = 'report.docx'
    else:
        file = args.file

    # 用于确认表头行的关键词，如果一行中包含了list中的任意一个内容，就被认为是表头
    if args.keywords == 'std':
        keyWords = ['HEAD', 'H1', 'A']
    else:
        keyWords = args.keywords

    # 确认表头时的行数限制，即只在前面limit行内找表头
    if args.limit == 'std':
        limit = 5
    else:
        limit = args.limit

    # 输出的文件名，默认为table_加原文件名
    if args.output == 'std':
        targetFile = f"result/table_{file.split('.')[0]}.xlsx"
    else:
        targetFile = args.output
    
    # 提取所有的表，保存在目标excel中，每一个sheet保存一个表
    extract_all_table(f"report/{file}", keyWords, limit, targetFile)
